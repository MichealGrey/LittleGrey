from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches

from src.core.types import WORD_TOOL_SCHEMA, ToolResult
from src.tools.base import BaseTool


class WordTool(BaseTool):
    name = "word_tool"
    description = "创建和编辑 Word 文档"
    input_schema = WORD_TOOL_SCHEMA
    is_risky = True

    def execute(self, params: dict[str, Any]) -> ToolResult:
        action = params["action"]
        try:
            if action == "create":
                return self._create(params)
            elif action == "read":
                return self._read(params)
            elif action == "write":
                return self._write(params)
            elif action == "add_table":
                return self._add_table(params)
            elif action == "add_image":
                return self._add_image(params)
            else:
                return ToolResult(success=False, message=f"未知操作: {action}")
        except Exception as e:
            return ToolResult(success=False, message=f"Word操作失败: {e}")

    def _create(self, params: dict[str, Any]) -> ToolResult:
        file_path = Path(params["file_path"])
        file_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        title = params.get("title")
        if title:
            doc.add_heading(title, level=0)

        paragraphs = params.get("paragraphs", [])
        for para in paragraphs:
            if isinstance(para, str):
                doc.add_paragraph(para)
            elif isinstance(para, dict):
                text = para.get("text", "")
                style = para.get("style", "Normal")
                if style.startswith("Heading"):
                    level = int(style.replace("Heading", "")) if style != "Heading" else 1
                    doc.add_heading(text, level=level)
                else:
                    doc.add_paragraph(text, style=style)

        doc.save(str(file_path))
        return ToolResult(
            success=True,
            file_path=str(file_path),
            message=f"Word文档已创建: {file_path}",
        )

    def _read(self, params: dict[str, Any]) -> ToolResult:
        file_path = Path(params["file_path"])
        if not file_path.exists():
            return ToolResult(success=False, message=f"文件不存在: {file_path}")

        doc = Document(str(file_path))
        paragraphs = []
        for para in doc.paragraphs:
            paragraphs.append({
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
            })

        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append(rows)

        return ToolResult(
            success=True,
            data={"paragraphs": paragraphs, "tables": tables},
            message=f"读取 {len(paragraphs)} 段落, {len(tables)} 表格",
        )

    def _write(self, params: dict[str, Any]) -> ToolResult:
        file_path = Path(params["file_path"])
        if not file_path.exists():
            return ToolResult(success=False, message=f"文件不存在: {file_path}")

        doc = Document(str(file_path))
        content = params.get("content", "")
        position = params.get("position")

        if position is not None and 0 <= position <= len(doc.paragraphs):
            if position == len(doc.paragraphs):
                doc.add_paragraph(content)
            else:
                doc.paragraphs[position].insert_text_before(content)
        else:
            doc.add_paragraph(content)

        doc.save(str(file_path))
        return ToolResult(
            success=True,
            file_path=str(file_path),
            message=f"Word文档已更新: {file_path}",
        )

    def _add_table(self, params: dict[str, Any]) -> ToolResult:
        file_path = Path(params["file_path"])
        if not file_path.exists():
            return ToolResult(success=False, message=f"文件不存在: {file_path}")

        doc = Document(str(file_path))
        table_data = params.get("table", {})
        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])

        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"

        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = str(header)

        for row_idx, row_data in enumerate(rows, 1):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(headers):
                    table.rows[row_idx].cells[col_idx].text = str(cell_data)

        doc.save(str(file_path))
        return ToolResult(
            success=True,
            file_path=str(file_path),
            message=f"表格已添加到 {file_path}",
        )

    def _add_image(self, params: dict[str, Any]) -> ToolResult:
        file_path = Path(params["file_path"])
        image_path = params.get("image_path", "")

        if not file_path.exists():
            return ToolResult(success=False, message=f"文件不存在: {file_path}")
        if not Path(image_path).exists():
            return ToolResult(success=False, message=f"图片不存在: {image_path}")

        doc = Document(str(file_path))
        doc.add_picture(image_path, width=Inches(5))
        doc.save(str(file_path))

        return ToolResult(
            success=True,
            file_path=str(file_path),
            message=f"图片已添加到 {file_path}",
        )

    def _risk_check(self, params: dict[str, Any]) -> tuple[bool, str]:
        action = params.get("action", "")
        file_path = params.get("file_path", "")
        if action in ("create", "write", "add_table", "add_image") and Path(file_path).exists():
            return True, f"文件 {file_path} 已存在，继续操作将修改原有内容"
        return False, ""
